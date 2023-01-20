from docx import Document
import json
def generate_docx_from_jsondict(jsondict='translated_dictionary.json',docx_name='translated_word.docx',mode='lesson'):
    with open(jsondict, 'r') as f:
        # Load the JSON data into a Python dictionary
        translated_dict = json.load(f)
    # Open the document
    document = Document()
    for key in translated_dict.keys():
        heading = document.add_heading(key, 2)
        # Add a new paragraph
        if mode == 'lesson':
            text=translated_dict[key]
            paragraph = document.add_paragraph()
            # Add some text to the paragraph
            for i in range(len(text.split('\n'))):
                run = paragraph.add_run(text.split('\n')[i])
                if i == 0 or text.split('\n')[i].startswith('"') :
                    run.font.bold = True
                if i != len(text.split('\n'))-1:
                    # Add a new line
                    run.add_break()
                # Add some more text to the same paragraph
                # run = paragraph.add_run("\n".join(text.split('\n')[1:]))
                run.font.name = "Helvetica"
        if mode == 'list':
            for text in translated_dict[key]:
                paragraph = document.add_paragraph()
                # Add some text to the paragraph
                for i in range(len(text.split('\n'))):
                    run = paragraph.add_run(text.split('\n')[i])
                    if i == 0 or text.split('\n')[i].startswith('"') :
                        run.font.bold = True
                    if i != len(text.split('\n'))-1:
                        # Add a new line
                        run.add_break()
                    # Add some more text to the same paragraph
                    # run = paragraph.add_run("\n".join(text.split('\n')[1:]))
                    run.font.name = "Helvetica"
            ##spacing
    # Save the document
    document.save(docx_name)
